"""cv_writer.py — Rewrite CVs to match job descriptions. Outputs LaTeX→PDF and/or DOCX→PDF."""

import re
import os
import json
import subprocess
import shutil
import tempfile
from pathlib import Path
from collections import Counter, defaultdict
from difflib import SequenceMatcher

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from jinja2 import Environment, BaseLoader

DATA_DIR = Path(__file__).parent / "data"
CV_TEXT_DIR = DATA_DIR / "cv_texts"
OUTPUT_DIR = Path(__file__).parent / "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

BASE_CV_MAP = {
    "Demo_User_CV_English": ["software", "developer", "engineer", "python", "javascript",
                              "full-stack", "backend", "cloud", "data", "machine learning",
                              "react", "node", "aws", "docker", "agile", "sql"],
}

ENRICHMENT_SOURCES = {
    "Demo_User_CV_English": [],
}


def load_cv_text(cv_key: str) -> str:
    from cv_matcher import CV_PROFILES
    profile_entry = CV_PROFILES.get(cv_key)
    if profile_entry:
        pdf_filename = profile_entry.get("file", "")
        txt_filename = Path(pdf_filename).stem + ".txt"
        path = CV_TEXT_DIR / txt_filename
        if path.exists():
            return path.read_text(encoding="utf-8")
    for f in CV_TEXT_DIR.iterdir():
        if cv_key.lower() in f.stem.lower().replace(" ", "_"):
            return f.read_text(encoding="utf-8")
    raise FileNotFoundError(f"No CV text found for '{cv_key}' in {CV_TEXT_DIR}")


def load_profile_json() -> dict:
    p = Path(__file__).parent / "profile_data.json"
    return json.loads(p.read_text(encoding="utf-8"))


SECTION_HEADERS = [
    "professional summary", "profile", "professional profile", "objective",
    "key skills", "skills", "core competencies", "additional skills",
    "professional experience", "experience", "work experience", "employment",
    "current position",
    "education", "academic background",
    "certifications", "courses", "training",
    "languages",
    "other information",
    "references",
]


def _guess_section(line: str) -> str | None:
    cleaned = re.sub(r"[_#*{}\[\]()]", "", line).strip().lower()
    for h in SECTION_HEADERS:
        if cleaned == h or cleaned.startswith(h + ":") or cleaned.startswith(h + " "):
            return h
    if line.strip().isupper() and len(line.strip()) > 3 and len(line.strip()) < 50:
        candidate = line.strip().lower().rstrip(":")
        if candidate in SECTION_HEADERS:
            return candidate
    return None


def parse_cv_text(text: str) -> dict:
    sections = {}
    current_section = "_header"
    sections[current_section] = []

    for line in text.split("\n"):
        header = _guess_section(line)
        if header:
            current_section = header
            sections.setdefault(current_section, [])
        elif line.strip():
            sections.setdefault(current_section, []).append(line.strip())
        else:
            sections.setdefault(current_section, []).append("")

    for k in sections:
        while sections[k] and sections[k][-1] == "":
            sections[k].pop()

    return sections


def parse_experience_block(lines: list[str]) -> list[dict]:
    roles = []
    current_role = None
    current_bullets = []

    for line in lines:
        if not line.strip():
            continue
        if current_role is None:
            current_role = {"title": line, "institution": "", "period": "", "bullets": []}
        elif current_role and not current_role["institution"] and not line.startswith("•") and not line.startswith("-") and not line.startswith("*") and not re.match(r"^\d", line):
            current_role["institution"] = line
        elif current_role and not current_role["period"] and re.match(r"^[A-Z][a-z]+\s+\d{4}", line):
            current_role["period"] = line
        elif line.startswith("•") or line.startswith("-") or line.startswith("*"):
            current_bullets.append(line.lstrip("•-* ").strip())
        elif re.match(r"^[A-Z][a-z]+\s+\d{4}\s*(–|-|to|–)\s*", line):
            current_role["period"] = line
        elif current_role and current_role["institution"] and not current_role["period"] and re.match(r"^[A-Za-z]", line):
            next_header = _guess_section(line)
            if next_header:
                if current_role:
                    current_role["bullets"] = current_bullets
                    roles.append(current_role)
                current_role = {"title": line, "institution": "", "period": "", "bullets": []}
                current_bullets = []
            elif len(line) < 80 and not line.endswith("."):
                current_role["institution"] += " " + line
            else:
                current_bullets.append(line)
        else:
            current_bullets.append(line)

    if current_role:
        current_role["bullets"] = current_bullets
        roles.append(current_role)

    return roles


def extract_keywords(text: str) -> list[str]:
    stopwords = {"the", "a", "an", "and", "or", "of", "in", "to", "for", "with",
                 "on", "at", "by", "is", "are", "was", "were", "be", "been",
                 "we", "you", "they", "it", "our", "your", "their", "its",
                 "this", "that", "these", "those", "from", "as", "per", "etc",
                 "using", "such", "including", "within", "about",
                 "between", "through", "during", "before", "after", "above",
                 "below", "up", "down", "out", "off", "over", "under", "again",
                 "further", "then", "once", "here", "there", "all", "each",
                 "every", "both", "few", "more", "most", "other", "some",
                 "no", "nor", "not", "only", "own", "same", "so", "than",
                 "too", "very", "just", "also", "well", "has", "have", "had",
                 "do", "does", "did", "done", "will", "would", "could",
                 "should", "may", "might", "shall", "can", "need"}
    words = re.findall(r"[a-zA-Z][a-zA-Z+.#\-]+", text.lower())
    return [w for w in words if w not in stopwords and len(w) > 2]


def extract_skills_from_text(text: str) -> list[str]:
    from skills_analyzer import ALL_SKILLS
    text_lower = text.lower()
    found = []
    for cat, skills in ALL_SKILLS.items():
        for sk in skills:
            if sk.lower() in text_lower and sk.lower() not in found:
                found.append(sk)
    return found


def score_bullet_relevance(bullet: str, keywords: list[str], skill_keywords: list[str]) -> float:
    bullet_lower = bullet.lower()
    kw_score = sum(2 for kw in keywords if kw in bullet_lower)
    skill_score = sum(3 for sk in skill_keywords if sk.lower() in bullet_lower)
    return kw_score + skill_score


def score_role_relevance(role_title: str, role_institution: str, role_bullets: list[str],
                         keywords: list[str], skill_keywords: list[str]) -> float:
    title_score = sum(3 for kw in keywords if kw in role_title.lower())
    inst_score = sum(1 for kw in keywords if kw in role_institution.lower())
    bullet_score = sum(score_bullet_relevance(b, keywords, skill_keywords) for b in role_bullets)
    return title_score + inst_score + bullet_score


def build_skill_keywords_from_job(job_title: str, job_desc: str) -> list[str]:
    from skills_analyzer import ALL_SKILLS
    text = f"{job_title} {job_desc}".lower()
    matched = []
    for cat, skills in ALL_SKILLS.items():
        for sk in skills:
            if sk.lower() in text:
                matched.append(sk)
    return matched


def rewrite_professional_summary(base_sections: dict, job_title: str, job_desc: str,
                                  keywords: list[str], profile: dict) -> str:
    original = ""
    for section in ["professional summary", "profile", "professional profile", "objective"]:
        if section in base_sections:
            original = " ".join(base_sections[section])
            break

    degree = profile.get("education", [{}])[0].get("degree", "Master's Degree in Law") if profile.get("education") else "Master's Degree in Law"
    job_keywords = ", ".join(keywords[:5])
    summary = (
        f"{degree} with extensive professional experience in legal administration, "
        f"sales management, and customer service. Skilled in {job_keywords}. "
        f"Proven track record of meeting targets, managing client relationships, "
        f"and handling complex administrative procedures. Seeking to apply expertise "
        f"to challenging opportunities in {job_title}."
    )
    return summary


def rewrite_skills_section(base_sections: dict, job_title: str, job_desc: str,
                            keywords: list[str], skill_keywords: list[str]) -> list[str]:
    original_skills = []
    for section in ["key skills", "skills", "core competencies", "additional skills"]:
        if section in base_sections:
            for line in base_sections[section]:
                original_skills.extend(re.split(r"[,;•\-\|]+", line))
            break

    original_skills = [s.strip() for s in original_skills if s.strip()]

    if not original_skills:
        return ["Legal procedures, Sales management, CRM, Microsoft Office, Customer Service"]

    scored = []
    for skill in original_skills:
        skill_lower = skill.lower()
        kw_score = sum(2 for kw in keywords if kw in skill_lower)
        sk_score = sum(3 for sk in skill_keywords if sk.lower() in skill_lower)
        scored.append((kw_score + sk_score, skill))

    scored.sort(key=lambda x: -x[0])

    result = []
    cat_skills = defaultdict(list)
    profile_json_data = load_profile_json()

    for score, skill in scored:
        categorized = False
        for cat_name, cat_skills_list in profile_json_data.get("technical_skills", {}).items():
            if skill.lower() in [s.lower() for s in cat_skills_list]:
                cat_skills[cat_name.replace("_", " ").title()].append(skill)
                categorized = True
                break
        if not categorized:
            cat_skills["General"].append(skill)

    for cat, skills in cat_skills.items():
        result.append(f"{cat}: {', '.join(skills)}")

    return result if result else original_skills[:3]


def rewrite_experience_section(base_sections: dict, job_title: str, job_desc: str,
                                keywords: list[str], skill_keywords: list[str],
                                enrichment_texts: list[str] = None) -> list[dict]:
    experience_text = []
    for section in ["professional experience", "experience", "work experience",
                    "employment", "current position"]:
        if section in base_sections:
            experience_text = base_sections[section]
            break

    if not experience_text:
        return []

    from profile import load_profile
    profile = load_profile()
    roles = []

    raw_roles = parse_experience_block(experience_text)

    if not raw_roles and profile.get("experience"):
        for exp in profile["experience"]:
            roles.append({
                "title": exp.get("role", ""),
                "institution": exp.get("institution", ""),
                "period": exp.get("period", ""),
                "bullets": exp.get("topics", []),
            })
    else:
        roles = raw_roles

    scored_roles = []
    for role in roles:
        score = score_role_relevance(
            role.get("title", ""),
            role.get("institution", ""),
            role.get("bullets", []),
            keywords,
            skill_keywords
        )
        scored_roles.append((score, role))

    scored_roles.sort(key=lambda x: -x[0])

    rewritten = []
    for score, role in scored_roles:
        bullets = role.get("bullets", [])
        if not bullets:
            rewritten.append(role)
            continue

        scored_bullets = []
        for b in bullets:
            bs = score_bullet_relevance(b, keywords, skill_keywords)
            scored_bullets.append((bs, b))

        scored_bullets.sort(key=lambda x: -x[0])

        kept_bullets = [b for s, b in scored_bullets if s > 0][:5]

        if enrichment_texts and len(kept_bullets) < 3:
            for et in enrichment_texts:
                for bullet in re.findall(r"(?:•|-|\*)\s*(.+?)(?:\n|$)", et):
                    bs = score_bullet_relevance(bullet, keywords, skill_keywords)
                    if bs > 0 and bullet not in kept_bullets:
                        kept_bullets.append(bullet)
                        if len(kept_bullets) >= 5:
                            break
                if len(kept_bullets) >= 5:
                    break

        if not kept_bullets and bullets:
            kept_bullets = bullets[:3]

        role["bullets"] = kept_bullets
        rewritten.append(role)

    return rewritten


def load_enrichment_texts(base_cv_key: str) -> list[str]:
    sources = ENRICHMENT_SOURCES.get(base_cv_key, [])
    texts = []
    for key in sources:
        try:
            texts.append(load_cv_text(key))
        except (FileNotFoundError, KeyError):
            pass
    return texts


def auto_select_base_cv(job_title: str, job_desc: str = "") -> str:
    text = f"{job_title} {job_desc}".lower()
    best_key = "Demo_User_CV_English"
    best_score = 0

    def kw_in_text(kw: str, txt: str) -> bool:
        kw = kw.lower()
        if kw in txt:
            return True
        words = kw.split()
        if len(words) > 1 and all(w in txt for w in words):
            return True
        return False

    for key, keywords in BASE_CV_MAP.items():
        score = sum(3 for kw in keywords if kw_in_text(kw, text))
        for kw in keywords:
            if kw_in_text(kw, job_title.lower()):
                score += 5
        if score > best_score:
            best_score = score
            best_key = key

    return best_key


LATEX_TEMPLATE = r"""\documentclass[11pt,a4paper]{article}

\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{geometry}
\geometry{margin=2cm, top=1.5cm, bottom=1.5cm}
\usepackage{hyperref}
\usepackage{titlesec}
\usepackage{enumitem}
\usepackage{xcolor}
\usepackage{parskip}

\definecolor{primary}{RGB}{30, 60, 120}
\definecolor{rulecolor}{RGB}{60, 100, 180}

\titleformat{\section}{\Large\bfseries\color{primary}}{}{0em}{}[\vspace{-0.3em}\color{rulecolor}\rule{\textwidth}{0.5pt}]
\titlespacing*{\section}{0pt}{1.2em}{0.6em}

\titleformat{\subsection}[hang]{\bfseries}{}{0em}{}
\titlespacing*{\subsection}{0pt}{0.8em}{0.2em}

\setlist[itemize]{leftmargin=1.2em, itemsep=0.1em, topsep=0.2em}

\renewcommand{\familydefault}{\sfdefault}

\begin{document}

\begin{center}
    {\LARGE\bfseries\color{primary} {{ name }} } \\[0.3em]
    {\small {{ contact_line }} } \\[0.2em]
    {\small \emph{{ contact_extra }} }
\end{center}

{% if summary %}
\section{Professional Summary}
{{ summary }}
{% endif %}

{% if skills %}
\section{Key Skills}
\begin{itemize}[nosep]
{% for skill_line in skills %}
    \item {{ skill_line }}
{% endfor %}
\end{itemize}
{% endif %}

{% if experience %}
\section{Professional Experience}
{% for role in experience %}
\subsection*{\textbf{{ role.title }}} \hfill \emph{{ role.period }}\\
\textit{{ role.institution }}
\begin{itemize}[nosep]
{% for bullet in role.bullets %}
    \item {{ bullet }}
{% endfor %}
\end{itemize}
{% endfor %}
{% endif %}

{% if education %}
\section{Education}
\begin{itemize}[nosep]
{% for entry in education %}
    \item \textbf{{ entry.degree }} — {{ entry.institution }} ({{ entry.year }})
{% endfor %}
\end{itemize}
{% endif %}

{% if languages %}
\section{Languages}
{{ languages }}
{% endif %}

\end{document}
"""


def _escape_latex(text: str) -> str:
    replacements = {
        '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#',
        '_': r'\_', '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}',
        '^': r'\^{}', '\\': r'\textbackslash{}',
        '—': r'---', '–': r'--', '•': r'$\bullet$ ',
        "'": r"'", "'": r"'", '"': r"``", '"': r"''",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text


def render_latex(sections: dict, output_path: str) -> str:
    env = Environment(loader=BaseLoader())
    template = env.from_string(LATEX_TEMPLATE)

    name = sections.get("name", "Demo User")
    contact_line = sections.get("contact_line", "")
    contact_extra = sections.get("contact_extra", "")
    summary = sections.get("summary", "")
    skills = sections.get("skills", [])
    experience = sections.get("experience", [])
    education = sections.get("education", [])
    languages = sections.get("languages", "")

    name = _escape_latex(name)
    contact_line = _escape_latex(contact_line)
    contact_extra = _escape_latex(contact_extra)
    summary = _escape_latex(summary)
    skills = [_escape_latex(s) for s in skills]
    languages = _escape_latex(languages)
    for role in experience:
        role["title"] = _escape_latex(role.get("title", ""))
        role["institution"] = _escape_latex(role.get("institution", ""))
        role["period"] = _escape_latex(role.get("period", ""))
        role["bullets"] = [_escape_latex(b) for b in role.get("bullets", [])]
    for entry in education:
        for k in entry:
            if isinstance(entry[k], str):
                entry[k] = _escape_latex(entry[k])

    latex_source = template.render(
        name=name,
        contact_line=contact_line,
        contact_extra=contact_extra,
        summary=summary,
        skills=skills,
        experience=experience,
        education=education,
        languages=languages,
    )

    tex_path = output_path + ".tex"
    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(latex_source)

    original_dir = os.getcwd()
    os.chdir(os.path.dirname(tex_path) or ".")
    try:
        result = subprocess.run(
            ["pdflatex", "-interaction=nonstopmode", os.path.basename(tex_path)],
            capture_output=True, timeout=60
        )
        subprocess.run(
            ["pdflatex", "-interaction=nonstopmode", os.path.basename(tex_path)],
            capture_output=True, timeout=60
        )
    finally:
        os.chdir(original_dir)

    pdf_path = output_path + ".pdf"
    if not os.path.exists(pdf_path):
        err = result.stderr.decode("latin-1", errors="replace")[:500] if result.stderr else "unknown"
        raise RuntimeError(f"pdflatex failed to produce PDF:\n{err}")

    for ext in [".aux", ".log", ".out"]:
        aux = output_path + ext
        if os.path.exists(aux):
            os.remove(aux)

    return latex_source, pdf_path


def render_docx(sections: dict, output_path: str) -> str:
    if not HAS_DOCX:
        raise ImportError("python-docx not installed. Run: pip3 install python-docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    name = sections.get("name", "Demo User")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(30, 60, 120)

    contact_line = sections.get("contact_line", "")
    if contact_line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact_line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

    contact_extra = sections.get("contact_extra", "")
    if contact_extra:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact_extra)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True

    def add_section_header(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(30, 60, 120)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def add_bullet(text: str):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)

    summary = sections.get("summary", "")
    if summary:
        add_section_header("Professional Summary")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(6)

    skills = sections.get("skills", [])
    if skills:
        add_section_header("Key Skills")
        for s in skills:
            add_bullet(s)

    experience = sections.get("experience", [])
    if experience:
        add_section_header("Professional Experience")
        for role in experience:
            p = doc.add_paragraph()
            run = p.add_run(role.get("title", ""))
            run.bold = True
            run.font.size = Pt(11)
            run2 = p.add_run(f"  —  {role.get('period', '')}")
            run2.italic = True
            run2.font.size = Pt(10)

            p = doc.add_paragraph()
            run = p.add_run(role.get("institution", ""))
            run.italic = True
            run.font.size = Pt(10)

            for bullet in role.get("bullets", []):
                add_bullet(bullet)

    education = sections.get("education", [])
    if education:
        add_section_header("Education")
        for entry in education:
            p = doc.add_paragraph()
            run = p.add_run(f"{entry.get('degree', '')} — {entry.get('institution', '')} ({entry.get('year', '')})")
            run.font.size = Pt(11)

    languages = sections.get("languages", "")
    if languages:
        add_section_header("Languages")
        p = doc.add_paragraph(languages)

    docx_path = output_path + ".docx"
    doc.save(docx_path)

    pdf_path = output_path + ".pdf"
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", os.path.dirname(docx_path) or ".", docx_path],
            capture_output=True, text=True, timeout=60,
            env={**os.environ, "HOME": os.path.expanduser("~")}
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice conversion timed out")
    except FileNotFoundError:
        raise RuntimeError("libreoffice not found. Install it for DOCX→PDF conversion.")

    if not os.path.exists(pdf_path):
        raise RuntimeError(f"LibreOffice failed to produce PDF from {docx_path}")

    return docx_path, pdf_path


def rewrite_cv(
    base_cv_key: str,
    job_title: str,
    job_description: str = "",
    organization: str = "",
    output_dir: str | None = None,
    formats: list[str] | None = None,
    extra_cv_keys: list[str] | None = None,
) -> dict:
    if formats is None:
        formats = ["latex", "docx"]
    if output_dir is None:
        output_dir = str(OUTPUT_DIR)

    os.makedirs(output_dir, exist_ok=True)

    profile = load_profile()
    profile_json = load_profile_json()
    keywords = extract_keywords(f"{job_title} {job_description}")
    skill_keywords = build_skill_keywords_from_job(job_title, job_description)

    base_text = load_cv_text(base_cv_key)
    base_sections = parse_cv_text(base_text)

    enrichment_texts = []
    if extra_cv_keys:
        for k in extra_cv_keys:
            try:
                enrichment_texts.append(load_cv_text(k))
            except (FileNotFoundError, KeyError):
                pass
    enrichment_texts += load_enrichment_texts(base_cv_key)

    contact_line = (
        f"{profile.get('location', '')} — "
        f"{profile.get('email', '')} — "
        f"{profile.get('phone', '')}"
    )

    summary = rewrite_professional_summary(
        base_sections, job_title, job_description, keywords, profile
    )

    skills = rewrite_skills_section(
        base_sections, job_title, job_description, keywords, skill_keywords
    )

    experience = rewrite_experience_section(
        base_sections, job_title, job_description, keywords, skill_keywords,
        enrichment_texts
    )

    education = []
    for e in profile.get("education", []):
        education.append({
            "degree": e.get("degree", ""),
            "institution": e.get("institution", ""),
            "year": e.get("year", ""),
        })

    languages_list = profile.get("languages", {})
    languages = ", ".join(f"{lang} ({prof})" for lang, prof in languages_list.items())

    sections = {
        "name": profile.get("name", "Demo User"),
        "contact_line": contact_line,
        "contact_extra": "",
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "education": education,
        "languages": languages,
    }

    output_base = os.path.join(output_dir, f"CV_{organization.replace(' ', '_')}_{job_title.replace(' ', '_')[:40]}")
    output_base = re.sub(r"[^\w/\-_.]", "_", output_base)

    result = {
        "sections": sections,
        "format_results": {},
    }

    if "latex" in formats:
        try:
            latex_source, latex_pdf = render_latex(sections, output_base)
            tex_path = output_base + ".tex"
            result["latex_source"] = latex_source
            result["latex_path"] = tex_path
            result["latex_pdf"] = latex_pdf
            result["format_results"]["latex"] = {"tex": tex_path, "pdf": latex_pdf}
        except Exception as e:
            result["format_results"]["latex"] = {"error": str(e)}

    if "docx" in formats:
        try:
            docx_path, docx_pdf = render_docx(sections, output_base)
            result["docx_path"] = docx_path
            result["docx_pdf"] = docx_pdf
            result["format_results"]["docx"] = {"docx": docx_path, "pdf": docx_pdf}
        except Exception as e:
            result["format_results"]["docx"] = {"error": str(e)}

    return result


def auto_rewrite_cv(
    job_title: str,
    job_description: str = "",
    organization: str = "",
    base_cv_key: str | None = None,
    **kwargs
) -> dict:
    if base_cv_key is None:
        base_cv_key = auto_select_base_cv(job_title, job_description)
    return rewrite_cv(base_cv_key, job_title, job_description, organization, **kwargs)


def list_available_cvs() -> list[str]:
    from cv_matcher import CV_PROFILES
    available = []
    for key in CV_PROFILES:
        txt_path = CV_TEXT_DIR / f"{key}.txt"
        if not txt_path.exists():
            for f in CV_TEXT_DIR.iterdir():
                if key.lower() in f.stem.lower():
                    txt_path = f
                    break
        if txt_path.exists():
            available.append(key)
    return available


def get_base_cv_for_job(job_title: str) -> str:
    return auto_select_base_cv(job_title)


from profile import load_profile
profile_json = load_profile_json()
